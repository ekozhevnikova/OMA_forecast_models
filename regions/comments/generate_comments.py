import pandas as pd
import pandas as pd
import pymorphy3 as pmrph
import docx
import numpy as np
from pathlib import Path
from io_data.dates import Dates_Operations
from io_data.operations import Dict_Operations, File



class Generate_Comments:
    def __init__(self, month_num, day_num):
        self.month_num = month_num
        self.day_num = day_num
        
    
    @staticmethod
    def get_right_city_name(city_name):
        """
        Generates the right spelling of difficult russian cities
        """
        if city_name == 'Ростове-на-Дону' or city_name == 'ростове-на-дону' or city_name == 'Ростове-на-дону' or city_name == 'Ростов-На-Дону':
            return 'Ростов-на-Дону', 'Ростове-на-Дону'
        elif city_name == 'Санкт-Петербурге' or city_name == 'санкт-петербурге' or city_name == 'Санкт-петербурге' or city_name == 'Санкт-Петербург':
            return 'Санкт-Петербург', 'Санкт - Петербурге'
        elif city_name == 'Нижний новгороде' or city_name == 'нижний новгороде' or city_name == 'Нижний новгороде' or city_name == 'Нижний Новгород':
            return 'Нижний Новгород' ,'Нижнем Новгороде'
        elif city_name == 'Великом Новгороде' or city_name == 'великом новгороде' or city_name == 'Великом новгороде' or city_name == 'Великом Новгород':
            return 'Великий Новогород' ,'Великом Новгороде'
        else: return city_name, city_name
        

    def get_reason_channels__and__res_data(self, data, data_api, column_fact_month, column_last_14_days, column_prev_14_days, column_prev_to_fact_month, cond_grp, cond_share, cond_kus, cond_ttv):
        """
        data - DataFrame with Previous data, Current data and Differencies between them.
        data_api - data from API
        """
        reason_channels = {
        'kus+': [],
        'ttv+': [],
        'share': [],
        'share4': [],
        'kus-': [],
        'ttv-': [],
        }

        res_data_1 = pd.DataFrame(data = {'Телеканал': [], 'Город': [], 'Было': [], 'Стало': [], 'Изменение': []})
        res_data_2 = pd.DataFrame(data = {'Телеканал': [], 'Город': [], 'Было': [], 'Стало': [], 'Изменение': []})
        res_data_3 = pd.DataFrame(data = {'Телеканал': [], 'Город': [], 'Было': [], 'Стало': [], 'Изменение': []})
        res_data_4 = pd.DataFrame(data = {'Телеканал': [], 'Город': [], 'Было': [], 'Стало': [], 'Изменение': []})

        data_grp = data[data['Атрибут'] == 'GRP']
        data_volume = data[data['Атрибут'] == 'Объем']
        data_share = data[data['Атрибут'] == 'Share']
        data_kus = data[data['Атрибут'] == 'КУС']
        data_ttv = data[data['Атрибут'] == 'TTV']


        month_changes = ['Январь', 'Февраль', 'Март', 
                             'Апрель', 'Май', 'Июнь', 
                             'Июль', 'Август', 'Сентябрь', 
                             'Октябрь', 'Ноябрь', 'Декабрь']

        current_month = month_changes[self.month_num - 1]
        previous_month = month_changes[self.month_num - 2]
        month_before_previous = month_changes[self.month_num - 3]

        if self.day_num < 8:
            current_month = month_changes[self.month_num - 2]
            previous_month = month_changes[self.month_num - 3]
            month_before_previous = month_changes[self.month_num - 4]
            month__before_before_previous = month_changes[self.month_num - 5]
        else:
            current_month = month_changes[self.month_num - 1]
            previous_month = month_changes[self.month_num - 2]
            month_before_previous = month_changes[self.month_num - 3]
            month__before_before_previous = month_changes[self.month_num - 4]

        not_found_cities_channels = set()
        for month in month_changes[self.month_num - 1:]: 
            for irow, row in data_grp.iterrows():
                row_vol = data_volume[(data_volume['Телеканал'] == row['Телеканал']) & (data_volume['Город'] == row['Город'])]

                if len(row_vol) > 0:
                    row_vol = row_vol.iloc[0]
                else:
                    print(f'Не нашелся объем для Город = {row["Город"]} и Телеканал = {row["Телеканал"]}')
                    continue
                if abs(row[f'{month}.2']) >= cond_grp and (abs(row[f'{month}.2'] - row_vol[f'{month}.2'])) >= cond_grp:
                #if (abs(row[f'{month}.2']) - abs(row_vol[f'{month}.2'])) >= cond_grp:
                    #is_kus_4 = False
                    #is_ttv_4 = False
                    is_share_4 = False
                    
                    if len(res_data_1[(res_data_1['Телеканал'] == row['Телеканал']) & (res_data_1['Город'] == Generate_Comments.get_right_city_name(city_name = row['Город'].lower().title())[0])]) > 0:
                        continue
                    elif len(res_data_2[(res_data_2['Телеканал'] == row['Телеканал']) & (res_data_2['Город'] == Generate_Comments.get_right_city_name(city_name = row['Город'].lower().title())[0])]) > 0:
                        continue
                    elif len(res_data_3[(res_data_3['Телеканал'] == row['Телеканал']) & (res_data_3['Город'] == Generate_Comments.get_right_city_name(city_name = row['Город'].lower().title())[0])]) > 0:
                        continue
                    elif len(res_data_4[(res_data_4['Телеканал'] == row['Телеканал']) & (res_data_4['Город'] == Generate_Comments.get_right_city_name(city_name = row['Город'].lower().title())[0])]) > 0:
                        continue

                    row_share = data_share[(data_share['Телеканал'] == row['Телеканал']) & (data_share['Город'] == row['Город'])].iloc[0]
                    row_kus = data_kus[(data_kus['Телеканал'] == row['Телеканал']) & (data_kus['Город'] == row['Город'])].iloc[0]
                    row_ttv = data_ttv[(data_ttv['Телеканал'] == row['Телеканал']) & (data_ttv['Город'] == row['Город'])].iloc[0]
                    #print(row['Телеканал'], row['Город'])
                    row_share_api = data_api[(data_api['Телеканал'] == row['Телеканал']) & (data_api['Город'] == row['Город'])].iloc[0]
                    #print((data_api['Телеканал'] == row['Телеканал'], data_api['Город'] == row['Город']))

                    if ((row[f'{month}.2'] - row_vol[f'{month}.2']) * ((row_share_api[column_fact_month] / row_share[f'{previous_month}.1']) - 1)) > 0 and abs((row_share_api[column_fact_month] / row_share[f'{previous_month}.1']) - 1) > cond_share:
                        #print(((row[f'{month}.2'] - row_vol[f'{month}.2']) * ((row_share_api[column_fact_month] / row_share[f'{previous_month}.1']) - 1)) > 0 and abs((row_share_api[column_fact_month] / row_share[f'{previous_month}.1']) - 1) > cond_share)
                        res_data_1.loc[len(res_data_1.index)] = [row['Телеканал'], 
                                                             Generate_Comments.get_right_city_name(city_name = row['Город'].lower().title())[0], 
                                                             row_share[f'{previous_month}.1'], 
                                                             row_share_api[column_fact_month], 
                                                             ((row_share_api[column_fact_month] / row_share[f'{previous_month}.1']) - 1) * 100]


                    elif ((row[f'{month}.2'] - row_vol[f'{month}.2']) * ((row_share_api[column_fact_month] / ((row_share[f'{month_before_previous}.1'] + row_share[f'{previous_month}.1']) / 2.0)) - 1)) > 0 and \
                    abs((row_share_api[column_fact_month] / ((row_share[f'{month_before_previous}.1'] + row_share[f'{previous_month}.1']) / 2.0)) - 1) > cond_share:
                        res_data_2.loc[len(res_data_2.index)] = [row['Телеканал'], 
                                                             Generate_Comments.get_right_city_name(city_name = row['Город'].lower().title())[0], 
                                                             ((row_share[f'{month_before_previous}.1'] + row_share[f'{previous_month}.1']) / 2.0), 
                                                             row_share_api[column_fact_month], 
                                                             ((row_share_api[column_fact_month] / ((row_share[f'{month_before_previous}.1'] + row_share[f'{previous_month}.1']) / 2.0)) - 1) * 100]


                    elif (((row[f'{month}.2'] - row_vol[f'{month}.2']) * ((row_share_api[column_last_14_days] / row_share_api[column_prev_14_days]) - 1))) > 0 and abs((row_share_api[column_last_14_days] / row_share_api[column_prev_14_days]) - 1) > cond_share:

                        res_data_3.loc[len(res_data_3.index)] = [row['Телеканал'], 
                                                             Generate_Comments.get_right_city_name(city_name = row['Город'].lower().title())[0], 
                                                             row_share_api[column_prev_14_days], 
                                                             row_share_api[column_last_14_days], 
                                                             ((row_share_api[column_last_14_days] / row_share_api[column_prev_14_days]) - 1) * 100]
                        
                        
                    elif ((row[f'{month}.2'] - row_vol[f'{month}.2']) * ((row_share_api[column_prev_to_fact_month] / ((row_share[f'{month__before_before_previous}.1'] + row_share[f'{month_before_previous}.1']) / 2.0)) - 1)) > 0 and \
                    abs((row_share_api[column_prev_to_fact_month] / ((row_share[f'{month__before_before_previous}.1'] + row_share[f'{month_before_previous}.1']) / 2.0)) - 1) > cond_share:
                        is_share_4 = True
                        res_data_4.loc[len(res_data_4.index)] = [row['Телеканал'], 
                                                             Generate_Comments.get_right_city_name(city_name = row['Город'].lower().title())[0], 
                                                             ((row_share[f'{month__before_before_previous}.1'] + row_share[f'{month_before_previous}.1']) / 2.0), 
                                                             row_share_api[column_prev_to_fact_month], 
                                                             ((row_share_api[column_prev_to_fact_month] / ((row_share[f'{month__before_before_previous}.1'] + row_share[f'{month_before_previous}.1']) / 2.0)) - 1) * 100]
                        
                    
                    else:
                        not_found_cities_channels.add((Generate_Comments.get_right_city_name(city_name = row['Город'].lower().title())[0], row['Телеканал']))

                    is_grp_change_negative = row[f'{month}.2'] < 0

                    if abs(row_kus[f'{month}.2'] * 100) >= cond_kus:
                        is_need_to_take_share = False
                        is_kus_change_negative = row_kus[f'{month}.2'] < 0

                        if is_grp_change_negative and is_kus_change_negative:
                            reason_channels['kus-'].append([row['Телеканал'], row['Город']])
                        elif (not is_grp_change_negative) and (not is_kus_change_negative):
                            reason_channels['kus+'].append([row['Телеканал'], row['Город']])
                    if abs(row_ttv[f'{month}.2'] * 100) >= cond_ttv:
                        is_need_to_take_share = False
                        is_ttv_change_negative = row_ttv[f'{month}.2'] < 0

                        if is_grp_change_negative and is_ttv_change_negative:
                            reason_channels['ttv-'].append([row['Телеканал'], row['Город']])
                        elif (not is_grp_change_negative) and (not is_ttv_change_negative):
                            reason_channels['ttv+'].append([row['Телеканал'], row['Город']])
                    if not is_share_4:
                        reason_channels['share'].append([row['Телеканал'], row['Город']])
                    else:
                        reason_channels['share4'].append([row['Телеканал'], row['Город']])
        print('\033[1m' + 'Не удалось найти релевантый период для доли при объяснении следующих пар (город, канал):\n\n' + '\n'.join(list(map(lambda x: str(x[0]) + ', ' + str(x[1]), not_found_cities_channels))))           
        return reason_channels, res_data_1, res_data_2, res_data_3, res_data_4
    
    
    @staticmethod
    def get_cities_formatted(dict_reasons):
        """
        dict_reasons: reason_channels
        """
        morph = pmrph.MorphAnalyzer()
        channels = []
        cities = []

        reason_channels_formatted = {
            'kus+': [],
            'ttv+': [], 
            'share': [],
            'share4': [],
            'kus-': [],
            'ttv-': [], 
        }

        for key, channels_cities in dict_reasons.items():
            for channel_city in channels_cities:
                channel_name = channel_city[0]
                city_name = morph.parse(channel_city[1].strip())[0]
                reason_channels_formatted[key].append([channel_name, 
                                                       city_name.inflect({'loct'}).word.capitalize()])
        #correction of mistakes of spelling in some cities         
        for key, values in reason_channels_formatted.items():
            for i in range(len(values)):
                reason_channels_formatted[key][i][1] = Generate_Comments.get_right_city_name(city_name = values[i][1])[1]
        return reason_channels_formatted
    
    
    @staticmethod
    def get_cities_and_reasons(reason_channels_formatted):
        """
        Function returns a dict of dicts, where
        Key - channel name
        key - reasons (ttv+, ttv-, kus+, kus-, ttv_kus+, ttv_kus-, share)
        value - cities
        """
        cities = set()
        channels = set()

        for key, value in reason_channels_formatted.items():
            for channel_city in value:
                channels.add(channel_city[0])
                cities.add(channel_city[1])

        # channel -> phrase condition -> cities
        channel_reason_cities = dict()
        for channel in channels:
            channel_reason_cities[channel] = {
                'ttv_kus+': set(),
                'ttv+': set(),
                'kus+': set(),
                'share' : set(),
                'share4': set(),
                'ttv_kus-': set(),
                'ttv-': set(),
                'kus-': set(),
            }


        for channel in channels:
            for city in cities:
                if [channel, city] in reason_channels_formatted['ttv+'] and [channel, city] in reason_channels_formatted['kus+']:
                    channel_reason_cities[channel]['ttv_kus+'].add(city)
                if [channel, city] in reason_channels_formatted['ttv-'] and [channel, city] in reason_channels_formatted['kus-']:
                    channel_reason_cities[channel]['ttv_kus-'].add(city)
                if [channel, city] in reason_channels_formatted['ttv+']:
                    channel_reason_cities[channel]['ttv+'].add(city)
                if [channel, city] in reason_channels_formatted['ttv-']:
                    channel_reason_cities[channel]['ttv-'].add(city)
                if [channel, city] in reason_channels_formatted['kus+']:
                    channel_reason_cities[channel]['kus+'].add(city)
                if [channel, city] in reason_channels_formatted['kus-']:
                    channel_reason_cities[channel]['kus-'].add(city)
                if [channel, city] in reason_channels_formatted['share']:
                    channel_reason_cities[channel]['share'].add(city)
                elif [channel, city] in reason_channels_formatted['share4']:
                    channel_reason_cities[channel]['share4'].add(city)

        return channel_reason_cities
    
    
    @staticmethod
    def get_df_channel(dataframe, channel):
        """

        """
        df = dataframe[dataframe['Телеканал'] == channel]
        return df
    
    
    @staticmethod
    def rename_columns(dataframe, column_name_prev_period, column_name_period_now):
        """
            Rename columns in res_data DataFrame
        """
        dataframe_ = dataframe.rename(columns = {'Было': column_name_prev_period, 'Стало': column_name_period_now})
        return dataframe_
    
    @staticmethod
    def remove_duplicate_elements_in_set(set_1, set_2):
        """
        Removes or Saves duplicate elements in sets
        Args:
            set_1
            set_2
        Example:
            set1 = (1, 2, 3, 4, 5)
            set2 = (2, 4)
            res_unique = (1, 3, 5)
            res_duplicates = (2)
            
        """
        res_duplicates = set()
        res_unique = set()
        for i in set_1:
            if i in set_2:
                res_duplicates.add(i)
            if i not in set_2:
                res_unique.add(i)
        return res_duplicates, res_unique

    
    def get_explanations(self, channel_reason_cities: dict, sorted_keys: list, channels_group_1, channels_group_2, channels_group_3):
        """
        This function generate comments in a sepcific way.
        All channels are divided into 4 groups. For each group there is a set of comments.
        channel_reason_cities is a dict of dicts, where
        Key - is the name of the channel
        key - reasons (ttv+, ttv-, kus+, kus-, ttv_kus+, ttv_kus-, share)
        value - cities where there are reasons for keys
        """
        comment_header_1 = 'Изменение прогноза инвентаря в нижеперечисленных городах связано с динамикой доли канала в '
        comment_header_2 = 'Изменение прогноза инвентаря в нижеперечисленных городах связано со значительным изменением доли, начиная с '
        
        comment_start = ['Увеличение прогноза инвентаря в ', 
                         'Уменьшение прогноза инвентаря в ']
        
        comment_ttv_kus_plus = [' связано с ростом общего уровня телесмотрения, а также с ростом эффективности рекламной сетки.\n', 
                                  ' обусловлено ростом общего уровня телесмотрения и ростом эффективности рекламной сетки.\n']
        
        comment_ttv_kus_minus = [' связано с уменьшением общего уровня телесмотрения, а также с уменьшением эффективности рекламной сетки.\n', 
                                   ' обусловлено уменьшением общего уровня телесмотрения и уменьшением эффективности рекламной сетки.\n']
        
        comment_kus_plus = [' обусловлено ростом эффективности рекламной сетки.\n', 
                            ' дополнительно произошло за счет роста эффективности рекламной сетки.\n', 
                            ' произошло благодаря росту эффективности рекламной сетки.\n']
        
        comment_kus_minus = [' связано с уменьшением эффективности рекламной сетки.\n', 
                             ' дополнительно произошло за счет снижения эффективности рекламной сетки.\n', 
                             ' произошло из - за снижения эффективности рекламной сетки.\n']
        
        comment_ttv_plus = [' дополнительно связано с ростом общего уровня телесмотрения.\n', 
                            ' связано с высокими фактическими показателями телесмотрения. Тенденции проложены на будущий период.\n', 
                            ' обусловлено также ростом общего уровня телесмотрения.\n']
        
        comment_ttv_minus = [' связано с низкими фактическими показателями телесмотрения. Тенденции проложены на будущий период.\n', 
                             ' обусловлено снижением общего уровня телесмотрения.\n', 
                             ' связано с уменьшением общего уровня телесмотрения.\n']

        channel_phrase = dict()
        for channel, reason_cities in channel_reason_cities.items():
            channel_name = ''
            phrase_kus_ttv_plus = {
                'cities': set(),
                'phrase': ''
            }
            phrase_kus_ttv_minus = {
                'cities': set(),
                'phrase': ''
            }
            phrase_kus_plus = {
                'cities': set(),
                'phrase': ''
            }
            phrase_kus_minus = {
                'cities': set(),
                'phrase': ''
            }
            phrase_ttv_plus = {
                'cities': set(),
                'phrase': ''
            }
            phrase_ttv_minus = {
                'cities': set(),
                'phrase': ''
            }
            phrase_share = {
                'cities': set(),
                'phrase': ''
            }
            phrase_share4 = {
                'cities': set(),
                'phrase': ''
            }
            phrase_kus_ttv_plus_4 = {
                'cities': set(),
                'phrase': ''
            }
            phrase_kus_ttv_minus_4 = {
                'cities': set(),
                'phrase': ''
            }
            phrase_kus_plus_4 = {
                'cities': set(),
                'phrase': ''
            }
            phrase_kus_minus_4 = {
                'cities': set(),
                'phrase': ''
            }
            phrase_ttv_plus_4 = {
                'cities': set(),
                'phrase': ''
            }
            phrase_ttv_minus_4 = {
                'cities': set(),
                'phrase': ''
            }
            
            res_reason_cities_unique = {}
            res_reason_cities_duplicates = {}
            for reason, city in reason_cities.items():
                res_reason_cities_duplicates[reason] = Generate_Comments.remove_duplicate_elements_in_set(set_1 = reason_cities[reason], set_2 = reason_cities['share4'])[0] #для res_data_4
                res_reason_cities_unique[reason] = Generate_Comments.remove_duplicate_elements_in_set(set_1 = reason_cities[reason], set_2 = reason_cities['share4'])[1] #для всех случаев, кроме res_data_4
            
            #Комментарии, если в положительную сторону изменились и TTV, и КУС для всех случаев, кроме res_data_4
            if len(res_reason_cities_unique['ttv_kus+']) > 0:
                if channel_name in channels_group_1:
                    phrase_kus_ttv_plus = {
                        'cities': res_reason_cities_unique['ttv_kus+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_unique['ttv_kus+']) + comment_ttv_kus_plus[0]
                    }
                else:
                     phrase_kus_ttv_plus = {
                        'cities': res_reason_cities_unique['ttv_kus+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_unique['ttv_kus+']) + comment_ttv_kus_plus[1]
                    }
                    
            #Комментарии, если в положительную сторону изменились и TTV, и КУС для res_data_4
            if len(res_reason_cities_duplicates['ttv_kus+']) > 0:
                if channel_name in channels_group_1:
                    phrase_kus_ttv_plus_4 = {
                        'cities': res_reason_cities_duplicates['ttv_kus+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_duplicates['ttv_kus+']) + comment_ttv_kus_plus[0]
                    }
                else:
                    phrase_kus_ttv_plus_4 = {
                        'cities': res_reason_cities_duplicates['ttv_kus+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_duplicates['ttv_kus+']) + comment_ttv_kus_plus[1]
                    }
                    
            #Комментарии, если в отрицательную сторону изменились и TTV, и КУС для всех случаев, кроме res_data_4
            if len(res_reason_cities_unique['ttv_kus-']) > 0:
                if channel_name in channels_group_1:
                    phrase_kus_ttv_minus = {
                        'cities': res_reason_cities_unique['ttv_kus-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_unique['ttv_kus-']) + comment_ttv_kus_minus[0]
                    }
                else:
                    phrase_kus_ttv_minus = {
                        'cities': res_reason_cities_unique['ttv_kus-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_unique['ttv_kus-']) + comment_ttv_kus_minus[1]
                    }
                    
            #Комментарии, если в отрицательную сторону изменились и TTV, и КУС для res_data_4
            if len(res_reason_cities_duplicates['ttv_kus-']) > 0:
                if channel_name in channels_group_1:
                    phrase_kus_ttv_minus_4 = {
                        'cities': res_reason_cities_duplicates['ttv_kus-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_duplicates['ttv_kus-']) + comment_ttv_kus_minus[0]
                    }
                else:
                    phrase_kus_ttv_minus_4 = {
                        'cities': res_reason_cities_duplicates['ttv_kus-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_duplicates['ttv_kus-']) + comment_ttv_kus_minus[1]
                    }
            
            #Комментарии, если в положительную сторону изменился КУС для всех случаев, кроме res_data_4
            if len(res_reason_cities_unique['kus+']) > 0 and len(res_reason_cities_unique['ttv_kus+']) == 0:
                if channel_name in channels_group_2:
                    phrase_kus_plus = {
                        'cities': res_reason_cities_unique['kus+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_unique['kus+']) + comment_kus_plus[0]
                    }
                elif channel_name in channels_group_3:
                    phrase_kus_plus = {
                        'cities': res_reason_cities_unique['kus+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_unique['kus+']) + comment_kus_plus[1]
                    }
                else:
                    phrase_kus_plus = {
                        'cities': res_reason_cities_unique['kus+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_unique['kus+']) + comment_kus_plus[2]
                    }
            
            #Комментарии, если в положительную сторону изменился КУС для res_data_4
            if len(res_reason_cities_duplicates['kus+']) > 0 and len(res_reason_cities_duplicates['ttv_kus+']) == 0:
                if channel_name in channels_group_2:
                    phrase_kus_plus_4 = {
                        'cities': res_reason_cities_duplicates['kus+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_duplicates['kus+']) + comment_kus_plus[0]
                    }
                elif channel_name in channels_group_3:
                    phrase_kus_plus_4 = {
                        'cities': res_reason_cities_duplicates['kus+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_duplicates['kus+']) + comment_kus_plus[1]
                    }
                else:
                    phrase_kus_plus_4 = {
                        'cities': res_reason_cities_duplicates['kus+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_duplicates['kus+']) + comment_kus_plus[2]
                    }
            
            #Комментарии, если в отрицательную сторону изменился КУС для всех случаев, кроме res_data_4
            if len(res_reason_cities_unique['kus-']) > 0 and len(res_reason_cities_unique['ttv_kus-']) == 0:
                if channel_name in channels_group_2:
                    phrase_kus_minus = {
                        'cities': res_reason_cities_unique['kus-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_unique['kus-']) + comment_kus_minus[0]
                    }
                elif channel_name in channels_group_3:
                    phrase_kus_minus = {
                        'cities': res_reason_cities_unique['kus-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_unique['kus-']) + comment_kus_minus[1]
                    }
                else:
                    phrase_kus_minus = {
                        'cities': res_reason_cities_unique['kus-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_unique['kus-']) + comment_kus_minus[2]
                    }
                    
            #Комментарии, если в отрицательную сторону изменился КУС для res_data_4
            if len(res_reason_cities_duplicates['kus-']) > 0 and len(res_reason_cities_duplicates['ttv_kus-']) == 0:
                if channel_name in channels_group_2:
                    phrase_kus_minus_4 = {
                        'cities': res_reason_cities_duplicates['kus-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_duplicates['kus-']) + comment_kus_minus[0]
                    }
                elif channel_name in channels_group_3:
                    phrase_kus_minus_4 = {
                        'cities': res_reason_cities_duplicates['kus-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_duplicates['kus-']) + comment_kus_minus[1]
                    }
                else:
                    phrase_kus_minus_4 = {
                        'cities': res_reason_cities_duplicates['kus-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_duplicates['kus-']) + comment_kus_minus[2]
                    }
            
            #Комментарии, если в положительную сторону изменился TTV для всех случаев, кроме res_data_4
            if len(res_reason_cities_unique['ttv+']) > 0 and len(res_reason_cities_unique['ttv_kus+']) == 0:
                if channel_name in channels_group_2:
                    phrase_ttv_plus = {
                        'cities': res_reason_cities_unique['ttv+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_unique['ttv+']) + comment_ttv_plus[0]
                    }
                elif channel_name in channels_group_3:
                    phrase_ttv_plus = {
                        'cities': res_reason_cities_unique['ttv+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_unique['ttv+']) + comment_ttv_plus[1]
                    }
                else:
                    phrase_ttv_plus = {
                        'cities': res_reason_cities_unique['ttv+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_unique['ttv+']) + comment_ttv_plus[2]
                    }
            
            #Комментарии, если в положительную сторону изменился TTV для res_data_4
            if len(res_reason_cities_duplicates['ttv+']) > 0 and len(res_reason_cities_duplicates['ttv_kus+']) == 0:
                if channel_name in channels_group_2:
                    phrase_ttv_plus_4 = {
                        'cities': res_reason_cities_duplicates['ttv+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_duplicates['ttv+']) + comment_ttv_plus[0]
                    }
                elif channel_name in channels_group_3:
                    phrase_ttv_plus_4 = {
                        'cities': res_reason_cities_duplicates['ttv+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_duplicates['ttv+']) + comment_ttv_plus[1]
                    }
                else:
                    phrase_ttv_plus_4 = {
                        'cities': res_reason_cities_duplicates['ttv+'],
                        'phrase': comment_start[0] + ', '.join(res_reason_cities_duplicates['ttv+']) + comment_ttv_plus[2]
                    }
            
            #Комментарии, если в отрицательную сторону изменился TTV для всех случаев, кроме res_data_4
            if len(res_reason_cities_unique['ttv-']) > 0 and len(res_reason_cities_unique['ttv_kus-']) == 0:
                if channel_name in channels_group_2:
                    phrase_ttv_minus = {
                        'cities': res_reason_cities_unique['ttv-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_unique['ttv-']) + comment_ttv_minus[0]
                    }
                elif channel_name in channels_group_3:
                    phrase_ttv_minus = {
                        'cities': res_reason_cities_unique['ttv-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_unique['ttv-']) + comment_ttv_minus[1]
                    }
                else:
                    phrase_ttv_minus = {
                        'cities': res_reason_cities_unique['ttv-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_unique['ttv-']) + comment_ttv_minus[2]
                    }
                    
            #Комментарии, если в отрицательную сторону изменился TTV для res_data_4
            if len(res_reason_cities_duplicates['ttv-']) > 0 and len(res_reason_cities_duplicates['ttv_kus-']) == 0:
                if channel_name in channels_group_2:
                    phrase_ttv_minus_4 = {
                        'cities': res_reason_cities_duplicates['ttv-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_duplicates['ttv-']) + comment_ttv_minus[0]
                    }
                elif channel_name in channels_group_3:
                    phrase_ttv_minus_4 = {
                        'cities': res_reason_cities_duplicates['ttv-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_duplicates['ttv-']) + comment_ttv_minus[1]
                    }
                else:
                    phrase_ttv_minus_4 = {
                        'cities': res_reason_cities_duplicates['ttv-'],
                        'phrase': comment_start[1] + ', '.join(res_reason_cities_duplicates['ttv-']) + comment_ttv_minus[2]
                    }

            if len(res_reason_cities_unique['share']) > 0:
                phrase_share = {
                    'cities': res_reason_cities_unique['share'],
                    'phrase': comment_header_1 + Dates_Operations.get_month(self.month_num, 1, 'Предложный')[0] + ':'
                }
                
            #Комментарий для шапки таблицы. Для случая, если существует res_data_4    
            if len(reason_cities['share4']) > 0:
                if len(reason_cities['share4']) <= 3:
                    phrase_share4 = {
                        'cities': reason_cities['share4'],
                        'phrase': 'Изменение прогноза инвентаря в ' + ', '.join(reason_cities['share4']) + ' связано со значительным изменением доли, начиная с ' + Dates_Operations.get_month(self.month_num, 2, 'Родительный')[0] + ':'
                    }
                else:
                    phrase_share4 = {
                        'cities': reason_cities['share4'],
                        'phrase': comment_header_2 + Dates_Operations.get_month(self.month_num, 2, 'Родительный')[0] + ':'
                    }
                            
            list_of_phrases = [
                phrase_kus_ttv_plus,
                phrase_kus_plus,
                phrase_ttv_plus,
                phrase_kus_ttv_minus,
                phrase_kus_minus,
                phrase_ttv_minus,
                phrase_kus_ttv_plus_4,
                phrase_kus_plus_4,
                phrase_ttv_plus_4,
                phrase_kus_ttv_minus_4,
                phrase_kus_minus_4,
                phrase_ttv_minus_4
            ]
            #print(list_of_phrases[3])
            list_of_phrases_non_empty = list_of_phrases

            #for cities_phrase in list_of_phrases:
             #   if cities_phrase['phrase'] != '':
             #       list_of_phrases_non_empty.append(cities_phrase)
            channel_phrase[channel + '\n'] = [phrase_share, phrase_share4] + list_of_phrases_non_empty
        channel_phrase_new = Dict_Operations(channel_phrase).sort_keys_in_dict(sorted_keys)
        #print(channel_phrase_new['ПЕРВЫЙ КАНАЛ\n'])
        return channel_phrase_new
    
    
    @staticmethod
    def get__channel_df(channel_phrase, res_data, column_name_prev_period, column_name_period_now):
        """
            Generate table with differencies in current period and previous period for each channel and city.
        """
        channel_df = {}

        for channel, _ in channel_phrase.items():
            channel_normalized = channel[:-1]
            df = Generate_Comments.get_df_channel(dataframe = res_data, channel = channel_normalized)
            df = df.drop(['Телеканал', 'Изменение'], axis = 1)
            df = Generate_Comments.rename_columns(df, column_name_prev_period, column_name_period_now)
            df.insert(loc = 3, column = 'Динамика', value = (
                np.array(list(map(float, df.iloc[:, 2].values))) / np.array(list(map(float, df.iloc[:, 1].values)))) - 1
                      )
            df[[column_name_prev_period]] = df[[column_name_prev_period]].applymap(lambda x: "{0:.2f}".format(x))
            df[[column_name_period_now]] = df[[column_name_period_now]].applymap(lambda x: "{0:.2f}".format(x))
            df[['Динамика']] = df[['Динамика']].applymap(lambda x: "{0:.0f}%".format(x * 100))

            channel_df[channel_normalized] = df

        return channel_df
    

    def to_file(self,
                output_filename, 
                res_data_1, 
                res_data_2, 
                res_data_3, 
                res_data_4,
                channel_phrase, 
                column_fact_month, 
                column_last_14_days,
                column_prev_14_days,
                column_prev_to_fact_month):
        """
        Save to file our report 
        """
        doc = docx.Document()
        File.set_style_doc_file(doc, 1.0, 1.0, 1.5, 1.5)
        if self.day_num < 8:
            last_month = str('Доля ') + Dates_Operations.get_month(self.month_num, 3, 'Родительный')[1]
            avg_per_2_last_months = str('Доля ') + Dates_Operations.get_month(self.month_num, 4, 'Родительный')[1] + ' - ' + Dates_Operations.get_month(self.month_num, 3, 'Родительный')[1]
            avg_per_3_last_months = str('Доля ') + Dates_Operations.get_month(self.month_num, 5, 'Родительный')[1] + ' - ' + Dates_Operations.get_month(self.month_num, 4, 'Родительный')[1]
        else:
            last_month = str('Доля ') + Dates_Operations.get_month(self.month_num, 2, 'Родительный')[1]
            avg_per_2_last_months = str('Доля ') + Dates_Operations.get_month(self.month_num, 3, 'Родительный')[1] + ' - ' + Dates_Operations.get_month(self.month_num, 2, 'Родительный')[1]
            avg_per_3_last_months = str('Доля ') + Dates_Operations.get_month(self.month_num, 4, 'Родительный')[1] + ' - ' + Dates_Operations.get_month(self.month_num, 3, 'Родительный')[1]
        channel_df_var_1 = Generate_Comments.get__channel_df(channel_phrase, res_data_1, last_month, column_fact_month)
        channel_df_var_2 = Generate_Comments.get__channel_df(channel_phrase, res_data_2, avg_per_2_last_months, column_fact_month)
        channel_df_var_3 = Generate_Comments.get__channel_df(channel_phrase, res_data_3, column_prev_14_days, column_last_14_days)
        channel_df_var_4 = Generate_Comments.get__channel_df(channel_phrase, res_data_4, avg_per_3_last_months, column_prev_to_fact_month)
        channel_df = [
            {'Вариант 1': channel_df_var_1},
            {'Вариант 2': channel_df_var_2},
            {'Вариант 3': channel_df_var_3},
            {'Вариант 4': channel_df_var_4}
        ]
        
        for c, p in channel_phrase.items():
            empty_dfs_count = 0
            empty_dfs_count_not4 = 0
            for idf, name_df in enumerate(channel_df):
                name, df = list(name_df.items())[0]
                #print(name)
                df = df[c.removesuffix('\n')]

                if len(df) == 0:
                    empty_dfs_count += 1
                if name != 'Вариант 4' and len(df) == 0:
                    empty_dfs_count_not4 += 1

            if empty_dfs_count == len(channel_df):
                continue
                
            p1 = doc.add_paragraph()
            
            #название канала
            r1 = p1.add_run(c.removesuffix('\n'))
            r1.bold = True
            

            for idf, name_df in enumerate(channel_df):
                name, df = list(name_df.items())[0]
                df = df[c.removesuffix('\n')]
                
                if len(df) == 0:
                    continue
                #таблица для всех случаев, кроме res_data_4
                #if len(channel_df) - empty_dfs_count > 1 and idf != 0:
                    #r_tmp = p1.add_run('\n')
                
               #шапка для всех случаев, кроме res_data_4 (p[0] - pharse_share)
                if name == 'Вариант 1' and empty_dfs_count_not4 != 3:
                    if p[0]['phrase'] != '':
                        pp = doc.add_paragraph(p[0]['phrase'])
                    #случай для res_data_4
                if name == 'Вариант 4':
                    #шапка для res_data_4
                    if p[1]['phrase'] != '':
                        p_tmp = doc.add_paragraph(p[1]['phrase'])
                
                t = doc.add_table(rows = df.shape[0] + 1, cols = df.shape[1])
                t.style = 'Table Grid'

                first_row_cells = t.rows[0].cells
                columns = list(df.columns)
                for col_idx in range(len(columns)):
                    first_row_cells[col_idx].text = columns[col_idx]

                for cell in first_row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # add the rest of the data frame
                for i in range(df.shape[0]):
                    for j in range(df.shape[1]):
                        t.cell(i + 1,j).text = str(df.values[i, j])
                       
                is_need_to_write = True
                if name != 'Вариант 4' and idf != len(channel_df) - 1:
                    for i in range(idf + 1, len(channel_df)-1):
                        if not list(channel_df[i].values())[0][c.removesuffix('\n')].empty:
                            is_need_to_write = False
                            break
                        
                p2 = doc.add_paragraph()
                if name != 'Вариант 4' and is_need_to_write:
                    #0: шапки для всех случаев, кроме res_data_4; 1 - шапки для res_data_4. Начиная со 2 элемента начинаются комментарии
                    for phrase in p[2:8]:
                        if phrase['phrase'] != '':
                            p2.add_run('\n')
                            p2.add_run(phrase['phrase'])
                elif name == 'Вариант 4':
                    #0: шапки для res_data_4
                    for phrase in p[8:]:
                        if phrase['phrase'] != '':
                            p2.add_run('\n')
                            p2.add_run(phrase['phrase'])
            #p3 = doc.add_paragraph()
            #p3.add_run('\n')
        doc.save(output_filename)
    
    
    def get__comments(self, filename, output_filename, data, data_api, column_fact_month, column_last_14_days, column_prev_14_days, column_prev_to_fact_month, sorted_keys, cond_grp, cond_share, cond_kus, cond_ttv, channels_group_1, channels_group_2, channels_group_3):
        reason_channels, res_data_1, res_data_2, res_data_3, res_data_4 = self.get_reason_channels__and__res_data(data, data_api, column_fact_month, column_last_14_days, column_prev_14_days, column_prev_to_fact_month, cond_grp, cond_share, cond_kus, cond_ttv)
        reason_channels_formatted = Generate_Comments.get_cities_formatted(reason_channels)
        channel_reason_cities = Generate_Comments.get_cities_and_reasons(reason_channels_formatted = reason_channels_formatted)
        channel_phrase = self.get_explanations(channel_reason_cities, sorted_keys, channels_group_1, channels_group_2, channels_group_3)
        self.to_file(output_filename, res_data_1, res_data_2, res_data_3, res_data_4, channel_phrase, column_fact_month, column_last_14_days, column_prev_14_days, column_prev_to_fact_month)

